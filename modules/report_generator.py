"""
report_generator.py
Формирование аналитической справки по налогоплательщику в HTML и Word (docx).

Режимы стиля справки (см. ТЗ п.12.3):
  - inspector   — «для налогового инспектора»: подробно, все таблицы и риски
  - client      — «для клиента»: простым языком, минимум терминологии
  - management  — «для руководителя»: коротко, только выводы и суммы
  - court       — «для суда/возражения»: аккуратно, только факты и документы,
                  без предположений и уровней риска в свободной форме

Во всех режимах соблюдается требование ТЗ п.9: никаких категоричных
обвинительных формулировок, только факт / расчёт / риск / рекомендация,
привязанные к конкретным суммам, периодам и источникам.
"""

from __future__ import annotations

import os
from dataclasses import dataclass, field
from datetime import date

import pandas as pd

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:  # pragma: no cover
    Document = None

try:
    from jinja2 import Environment, FileSystemLoader, select_autoescape
except ImportError:  # pragma: no cover
    Environment = None

TEMPLATES_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "templates")

RISK_COLOR = {"высокий": "#e05555", "средний": "#e8b93a", "низкий": "#5aa15a"}

MODE_TITLES = {
    "inspector": "для налогового инспектора",
    "client": "для клиента",
    "management": "для руководителя",
    "court": "для суда / возражения",
}


@dataclass
class ReportContext:
    taxpayer_name: str
    taxpayer_bin: str
    years_label: str
    mode: str
    generated_at: str
    profile: dict
    yearly_profile: dict
    dynamics: list[dict]
    buyers_by_year: dict[int, list[dict]]
    suppliers_by_year: dict[int, list[dict]]
    vat_reconciliation: list[dict]
    income_reconciliation: list[dict]
    payroll_check: list[dict]
    findings_by_level: dict[str, list[dict]]
    documents_to_request: list[dict]
    summary_text: str
    conclusion_text: str
    data_gaps: list[str]
    watermark_text: str | None = None


def _fmt_amount(v) -> str:
    if v is None or (isinstance(v, float) and pd.isna(v)):
        return "н/д"
    try:
        return f"{v:,.0f}".replace(",", " ")
    except (TypeError, ValueError):
        # защита от jinja2.Undefined и прочих нечисловых значений, когда в
        # profile/yearly_profile отсутствует ожидаемый ключ целиком (не
        # просто None) — приложение никогда не должно падать (см. ТЗ п.1)
        return "н/д"


def _fmt_pct(v) -> str:
    if v is None or (isinstance(v, float) and pd.isna(v)):
        return "н/д"
    try:
        return f"{v:.1f}%"
    except (TypeError, ValueError):
        return "н/д"


def build_data_gaps(has_taxpayer: bool, has_fno100: bool, has_fno200: bool, has_fno300: bool,
                     has_esf_purchases: bool, has_esf_sales: bool) -> list[str]:
    gaps = []
    if not has_taxpayer:
        gaps.append("сведения по налогоплательщику (карточка НП / СУР КГД)")
    if not has_esf_sales:
        gaps.append("ЭСФ по реализации")
    if not has_esf_purchases:
        gaps.append("ЭСФ по приобретению")
    if not has_fno100:
        gaps.append("ФНО 100.00 (для сверки дохода с реализацией по ЭСФ)")
    if not has_fno300:
        gaps.append("ФНО 300.00 (для камерального контроля по НДС)")
    if not has_fno200:
        gaps.append("ФНО 200.00/200.01 (для проверки зарплатных налогов)")
    return gaps


def build_summary_text(taxpayer_name: str, taxpayer_bin: str, regime: str | None, vat_payer,
                        oked_name: str | None, total_sales: float, total_purchases: float,
                        findings_by_level: dict, data_gaps: list[str]) -> str:
    vat_text = "является плательщиком НДС" if vat_payer is True else (
        "не является плательщиком НДС" if vat_payer is False else "статус плательщика НДС не определён по загруженным данным"
    )
    regime_text = regime or "не указан в загруженных данных"
    oked_text = oked_name or "не определён по загруженным данным"
    n_high = len(findings_by_level.get("высокий", []))
    n_med = len(findings_by_level.get("средний", []))
    n_low = len(findings_by_level.get("низкий", []))

    parts = [
        f"Налогоплательщик {taxpayer_name} (БИН/ИИН {taxpayer_bin}), {vat_text}. "
        f"Налоговый режим: {regime_text}. Основной вид деятельности: {oked_text}.",
        f"По данным ЭСФ общий объём реализации за анализируемый период составляет "
        f"{_fmt_amount(total_sales)} тенге без НДС, объём приобретений — {_fmt_amount(total_purchases)} тенге без НДС.",
    ]
    if n_high or n_med or n_low:
        parts.append(
            f"По результатам сопоставления данных выявлено risk-находок: высокий уровень риска — {n_high}, "
            f"средний — {n_med}, низкий — {n_low}. Все находки требуют дополнительной проверки и не являются "
            f"окончательным выводом о нарушении налогового законодательства."
        )
    else:
        parts.append("По результатам сопоставления загруженных данных существенных расхождений не выявлено.")
    if data_gaps:
        parts.append(
            "Недостаточно данных для окончательного вывода по отдельным направлениям анализа. "
            "Необходимо загрузить: " + ", ".join(data_gaps) + "."
        )
    return " ".join(parts)


def build_conclusion_text(findings_by_level: dict, data_gaps: list[str]) -> str:
    n_high = len(findings_by_level.get("высокий", []))
    n_med = len(findings_by_level.get("средний", []))
    lines = []
    if n_high:
        lines.append(
            f"По загруженным данным выявлено {n_high} находок с высоким уровнем риска — по ним требуется "
            f"первоочередная дополнительная проверка и запрос подтверждающих документов у налогоплательщика "
            f"и/или контрагентов."
        )
    if n_med:
        lines.append(
            f"Выявлено {n_med} находок со средним уровнем риска, включая концентрацию оборота на отдельных "
            f"контрагентах и отдельные расхождения — требуется уточнение по документам."
        )
    if not n_high and not n_med:
        lines.append("По имеющимся данным существенных признаков риска не выявлено.")
    lines.append(
        "Все приведённые в справке суммы и выводы основаны исключительно на загруженных файлах и не являются "
        "окончательным заключением о нарушении налогового законодательства. Однозначные выводы о фиктивности "
        "сделок или уклонении от уплаты налогов возможны только по результатам истребования и анализа первичных "
        "документов (договоров, актов, накладных, СНТ, платёжных документов) и подтверждения фактического "
        "исполнения обязательств контрагентами."
    )
    if data_gaps:
        lines.append("Для более полного анализа рекомендуется дополнительно загрузить: " + ", ".join(data_gaps) + ".")
    lines.append(
        "Рекомендуемые следующие действия: запросить документы по контрагентам и периодам с выявленными "
        "расхождениями (см. раздел «Документы, которые нужно запросить»), провести встречные проверки по "
        "контрагентам с признаками риска, уточнить период отнесения НДС в зачёт по кварталам с расхождениями."
    )
    return " ".join(lines)


def render_html(ctx: ReportContext) -> str:
    if Environment is None:
        raise RuntimeError("jinja2 не установлен")
    env = Environment(
        loader=FileSystemLoader(TEMPLATES_DIR),
        autoescape=select_autoescape(["html"]),
    )
    env.filters["amount"] = _fmt_amount
    env.filters["pct"] = _fmt_pct
    template = env.get_template("analytical_report.html")
    return template.render(ctx=ctx, risk_color=RISK_COLOR, mode_title=MODE_TITLES.get(ctx.mode, ctx.mode))


def _add_table(doc, headers: list[str], rows: list[list], col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = str(h)
        for p in hdr_cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = "" if val is None else str(val)
    return table


def render_docx(ctx: ReportContext, output_path: str) -> str:
    if Document is None:
        raise RuntimeError("python-docx не установлен")
    doc = Document()

    if ctx.watermark_text:
        wm_para = doc.add_paragraph()
        wm_run = wm_para.add_run(ctx.watermark_text)
        wm_run.bold = True
        wm_run.font.color.rgb = RGBColor(0xC0, 0x30, 0x30)
        wm_run.font.size = Pt(13)

    title = doc.add_heading(
        f"Аналитическая справка по налогоплательщику {ctx.taxpayer_name}, "
        f"БИН/ИИН {ctx.taxpayer_bin}, за период {ctx.years_label}", level=1
    )
    doc.add_paragraph(f"Режим справки: {MODE_TITLES.get(ctx.mode, ctx.mode)}. Сформировано: {ctx.generated_at}.")

    doc.add_heading("1. Краткий вывод", level=2)
    doc.add_paragraph(ctx.summary_text)

    doc.add_heading("2. Данные по налогоплательщику", level=2)
    p = ctx.profile
    info_rows = [
        ["БИН/ИИН", ctx.taxpayer_bin],
        ["Наименование", ctx.taxpayer_name],
        ["Дата регистрации", str(p.get("date_registration") or "н/д")],
        ["ОКЭД", f"{p.get('oked') or 'н/д'} — {p.get('oked_name') or ''}"],
        ["Налоговый режим", p.get("tax_regime") or "н/д"],
        ["Плательщик НДС", "Да" if p.get("vat_payer") is True else ("Нет" if p.get("vat_payer") is False else "н/д")],
        ["Дата постановки на НДС", str(p.get("vat_reg_date") or "н/д")],
        ["Налоговая задолженность", _fmt_amount(p.get("tax_debt_total"))],
    ]
    _add_table(doc, ["Показатель", "Значение"], info_rows)

    if ctx.mode != "court" and ctx.yearly_profile:
        doc.add_heading("Численность, налоговые поступления и КНН по годам", level=3)
        rows = []
        for y, d in sorted(ctx.yearly_profile.items()):
            rows.append([
                y, d.get("avg_headcount"), _fmt_amount(d.get("tax_revenue")),
                d.get("knn_taxpayer"), d.get("knn_industry"),
            ])
        _add_table(doc, ["Год", "Численность", "Налоговые поступления", "КНН НП, %", "КНН отрасли, %"], rows)

    doc.add_heading("3. Динамика по годам", level=2)
    rows = [[
        d["year"], _fmt_amount(d["sales_no_vat"]), _fmt_amount(d["sales_vat"]),
        _fmt_amount(d["purchases_no_vat"]), _fmt_amount(d["purchases_vat"]),
        _fmt_amount(d["diff_sales_purchases"]), (f"{d['tax_burden_pct']}%" if d.get("tax_burden_pct") else "н/д"),
    ] for d in ctx.dynamics]
    _add_table(doc, ["Год", "Реализация без НДС", "НДС реализация", "Приобретение без НДС",
                      "НДС приобретение", "Разница", "Налоговая нагрузка"], rows)

    if ctx.mode in ("inspector", "management"):
        doc.add_heading("4. Покупатели по годам (топ)", level=2)
        for y, items in sorted(ctx.buyers_by_year.items()):
            doc.add_heading(f"{y} год", level=3)
            rows = [[b["name"] or b["bin"], b["bin"], b["esf_count"], _fmt_amount(b["amount_no_vat"]),
                     _fmt_pct(b["share_pct"]), b.get("sector")] for b in items[:10]]
            _add_table(doc, ["Покупатель", "БИН", "Кол-во ЭСФ", "Сумма без НДС", "Доля", "Сектор"], rows)

        doc.add_heading("5. Поставщики по годам (топ)", level=2)
        for y, items in sorted(ctx.suppliers_by_year.items()):
            doc.add_heading(f"{y} год", level=3)
            rows = [[s["name"] or s["bin"], s["bin"], s["esf_count"], _fmt_amount(s["amount_no_vat"]),
                     _fmt_pct(s["share_pct"]), s.get("concentration_flag")] for s in items[:10]]
            _add_table(doc, ["Поставщик", "БИН", "Кол-во ЭСФ", "Сумма без НДС", "Доля", "Концентрация"], rows)

    if ctx.mode in ("inspector",):
        doc.add_heading("6. Сопоставление с ФНО", level=2)
        doc.add_heading("НДС: ЭСФ vs ФНО 300.00", level=3)
        rows = [[f"{v['year']} Q{v['quarter']}", _fmt_amount(v["esf_vat_charged"]), _fmt_amount(v.get("fno_vat_charged")),
                 _fmt_amount(v.get("diff_vat_charged")), _fmt_amount(v["esf_vat_credit"]),
                 _fmt_amount(v.get("fno_vat_credit")), _fmt_amount(v.get("diff_vat_credit"))]
                for v in ctx.vat_reconciliation]
        _add_table(doc, ["Период", "НДС начисл. (ЭСФ)", "НДС начисл. (ФНО)", "Разница",
                          "НДС в зачёт (ЭСФ)", "НДС в зачёт (ФНО)", "Разница"], rows)

        doc.add_heading("Доход: ЭСФ реализация vs ФНО 100.00", level=3)
        rows = [[i["year"], _fmt_amount(i["esf_sales_amount"]), _fmt_amount(i.get("fno_income")), _fmt_amount(i.get("diff"))]
                for i in ctx.income_reconciliation]
        _add_table(doc, ["Год", "Реализация (ЭСФ)", "Доход (ФНО 100.00)", "Разница"], rows)

    doc.add_heading("7. Камеральный контроль и риски", level=2)
    for level in ("высокий", "средний", "низкий"):
        items = ctx.findings_by_level.get(level, [])
        if not items:
            continue
        doc.add_heading(f"Уровень риска: {level}", level=3)
        for f in items:
            para = doc.add_paragraph()
            run = para.add_run(f"[{f['period']}] {f['risk_type']}. ")
            run.bold = True
            para.add_run(f"{f['description']} Что проверить: {f['what_to_check']} "
                         f"Возможные последствия: {f['possible_consequence']}")

    if ctx.mode in ("inspector", "client"):
        doc.add_heading("8. Документы, которые нужно запросить", level=2)
        rows = [[d["counterparty_or_period"], d["document"], d["reason"]] for d in ctx.documents_to_request]
        if rows:
            _add_table(doc, ["Контрагент / период", "Документ", "Основание"], rows)
        else:
            doc.add_paragraph("Дополнительные документы по выявленным находкам не требуются.")

    doc.add_heading("9. Предварительный вывод", level=2)
    doc.add_paragraph(ctx.conclusion_text)

    if ctx.watermark_text:
        doc.add_paragraph()
        wm_para2 = doc.add_paragraph()
        wm_run2 = wm_para2.add_run(ctx.watermark_text)
        wm_run2.bold = True
        wm_run2.font.color.rgb = RGBColor(0xC0, 0x30, 0x30)

    doc.save(output_path)
    return output_path
